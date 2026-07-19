"""
Generates docs/KFall_Pattern_Analysis_Report.docx -- the formal report companion
to docs/pattern_analysis.md. Run from the KFall project root:

    python3 docs/build_report.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(ROOT, "docs", "KFall_Pattern_Analysis_Report.docx")

HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)
BODY_FONT = "Calibri"


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = BODY_FONT
    run.font.color.rgb = HEADING_COLOR
    return h


def add_para(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    if align:
        p.alignment = align
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = BODY_FONT


def add_figure(doc, rel_path, caption, width_in=6.0):
    path = os.path.join(ROOT, rel_path)
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = add_para(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.runs[0].font.size = Pt(9.5)


def shade_cell(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = BODY_FONT
        shade_cell(hdr_cells[i], "1F3864")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = BODY_FONT
                    r.font.size = Pt(10.5)
    return table


def build():
    doc = Document()
    set_base_style(doc)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KFall Pre-Impact Fall Signal Pattern:\nValidation and Generalization Analysis")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = HEADING_COLOR

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("A signal-processing assessment of whether the pre-impact fall\n"
                        "signature in the KFall dataset generalizes across subjects and fall types")
    srun.italic = True
    srun.font.size = Pt(13)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Dataset: KFall (Yu, Jang & Xiong, 2021)\n"
                  "Scope: 32 subjects, 15 simulated fall types, 21 activities of daily living\n"
                  "Sensor: single 9-axis wearable IMU (accelerometer, gyroscope, orientation), 100 Hz").font.size = Pt(10.5)
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "This report documents an end-to-end signal-processing assessment of the pre-impact fall "
        "signature present in the KFall dataset. The objective was to determine whether the "
        "characteristic fall signal -- an acceleration dip followed by an impact spike -- is a "
        "consistent, generalizable pattern across the dataset's 32 subjects and 15 simulated fall "
        "types, or whether it is specific to individual subjects or fall types and therefore not a "
        "reliable basis for pre-impact fall detection.")
    add_para(doc,
        "The analysis proceeded in three stages: (1) validation of the detection logic against the "
        "dataset's originating publication, (2) sensitivity and specificity testing against both fall "
        "and non-fall (activities of daily living) recordings, and (3) direct visualization of signal "
        "shape using event-locked (phase-aligned) averaging across all fall trials.")
    add_para(doc,
        "Findings confirm that three of the four available signal channels -- acceleration magnitude, "
        "vertical velocity, and gyroscope magnitude -- exhibit a consistent, sharply-timed pattern that "
        "generalizes across the full dataset, verified at the level of individual trials. The fourth "
        "channel, orientation tilt, is comparatively weak and noisy as a standalone marker. Single "
        "fixed-threshold rules on these signals, while capturing the pattern reliably in fall trials, "
        "also trigger on a substantial share of ordinary daily activities and are therefore not "
        "sufficient as a standalone detector; a trained classifier, consistent with the originating "
        "publication's own benchmark results, is required for practical deployment.")

    # 2. Objective
    add_heading(doc, "2. Objective", level=1)
    add_para(doc,
        "To determine, using both inter-subject analysis (holding fall type constant, varying subject) "
        "and intra-subject analysis (holding subject constant, varying fall type), whether the "
        "pre-impact fall signal pattern observed in individual example trials is a global property of "
        "the KFall dataset, and to identify which specific sensor signals carry that pattern reliably.")

    # 3. Dataset Description
    add_heading(doc, "3. Dataset Description", level=1)
    add_para(doc,
        "KFall (Yu, Jang & Xiong, 2021, Frontiers in Aging Neuroscience) was developed from 32 young "
        "Korean adults wearing a nine-axis inertial sensor (accelerometer, gyroscope, and fused Euler "
        "orientation) on the lower back, sampled at 100 Hz. Each subject performed 21 types of "
        "activities of daily living (ADLs, task codes D01-D21) and 15 types of simulated falls (task "
        "codes F01-F15, internally referenced as task IDs 20-34), each labeled with a fall onset frame "
        "and a fall impact frame derived from synchronized video review. This analysis used all "
        "available recordings: 2,319 fall trials and 2,717 ADL trials across the 32 subjects.")

    # 4. Methodology
    add_heading(doc, "4. Methodology", level=1)

    add_heading(doc, "4.1 Signal Definitions", level=2)
    add_bullets(doc, [
        "ACC_M (acceleration magnitude): the Euclidean norm of the three-axis accelerometer reading, "
        "low-pass filtered at 5 Hz.",
        "VV (vertical velocity): body-frame acceleration rotated into the world frame using the "
        "sensor's fused orientation estimate, projected onto the vertical axis, and integrated over a "
        "bounded one-second window to avoid long-term drift.",
        "Tilt: the maximum of the absolute pitch and roll angles.",
        "Gyroscope magnitude: the Euclidean norm of the three-axis gyroscope reading, low-pass "
        "filtered at 5 Hz.",
    ])

    add_heading(doc, "4.2 Literature-Grounded Detection Rules", level=2)
    add_para(doc,
        "Rather than adopt an arbitrary threshold, the detection logic was grounded in the dataset's "
        "originating publication, which specifies a threshold-based benchmark algorithm combining "
        "acceleration magnitude, vertical velocity, and orientation. Three rules were defined and "
        "evaluated identically across all trials:")
    add_table(doc,
        ["Rule", "Definition", "Basis"],
        [
            ["A", "ACC_M < 0.8 g at any point in the trial", "Acceleration-only baseline"],
            ["B", "ACC_M < 0.8 g AND VV > 0.3 m/s, confirmed by tilt > 25 deg "
                  "within 10 subsequent frames", "Source publication's full algorithm"],
            ["C", "Peak gyroscope magnitude exceeds three times the subject's "
                  "standing baseline (fainting-while-sitting fall subset only)", "Source publication's "
                  "noted alternative sensitive axis for low-dynamic falls"],
        ])

    add_heading(doc, "4.3 Signal Validation", level=2)
    add_para(doc,
        "The vertical velocity computation was validated against a physical expectation: during a "
        "30-second quiet-standing recording, vertical velocity should remain near zero. Two "
        "implementation errors were identified and corrected through this check -- unbounded "
        "integration drift, and a sign convention inconsistent with the intended \"downward velocity\" "
        "interpretation. Following correction, Rule B's aggregate sensitivity (95.6%) closely matched "
        "the source publication's independently reported figure (95.50%), providing external "
        "confirmation that the corrected implementation is sound.")

    add_heading(doc, "4.4 Event-Locked Phase-Aligned Visualization", level=2)
    add_para(doc,
        "Because individual fall trials vary substantially in duration (approximately 70 frames for a "
        "brief sitting-and-fainting fall to over 500 frames for a fall while walking), signals were "
        "time-normalized onto a common phase axis prior to visualization: 0% denotes the labeled onset "
        "frame, 100% the labeled impact frame, extended to -50% through +150% for surrounding context. "
        "This is a standard event-locked averaging technique that allows trials of heterogeneous "
        "duration to be meaningfully overlaid and averaged. Two parallel analyses were produced: a "
        "three-signal version (acceleration magnitude, vertical velocity, tilt) and a four-signal "
        "version additionally incorporating gyroscope magnitude, maintained in separate subfolders "
        "(analysis_3panel/ and analysis_4panel/) for clarity.")

    # 5. Results
    add_heading(doc, "5. Results", level=1)

    add_heading(doc, "5.1 Sensitivity and Specificity", level=2)
    add_table(doc,
        ["Rule", "Sensitivity", "Specificity"],
        [
            ["A (acceleration only)", "100.0%", "17.4%"],
            ["B (acceleration + vertical velocity + tilt)", "95.6%", "34.9%"],
            ["Source publication's Threshold algorithm (reference)", "95.50%", "83.43%"],
        ])
    add_para(doc,
        "Acceleration magnitude alone (Rule A) detects every fall trial but also triggers on 82.6% of "
        "ADL trials, indicating negligible standalone discriminating power. Adding vertical velocity "
        "and tilt confirmation (Rule B) roughly doubles specificity for a marginal reduction in "
        "sensitivity, confirming that the combination of signals specified in the source publication is "
        "necessary rather than incidental. The residual gap to the source publication's reported "
        "specificity is attributed to differences in evaluation protocol: this analysis records whether "
        "a rule fires at any point across an entire recording, a stricter test than a streaming, "
        "debounced real-time evaluation.")
    add_figure(doc, "paper_threshold_validation/plots/sensitivity_specificity_comparison.png",
               "Figure 1. Sensitivity and specificity, Rule A versus Rule B.")

    add_heading(doc, "5.2 Activity-of-Daily-Living False-Positive Analysis", level=2)
    add_para(doc,
        "False-positive rates by ADL task show that dynamic activities -- jogging, gentle jumping, "
        "stumbling, quick stair descent, and rapid sit-to-stand transitions -- are the primary source "
        "of false triggers for both rules, reflecting genuine physical similarity to fall dynamics "
        "rather than a deficiency in threshold calibration. Rule B substantially reduces false-positive "
        "rates on several slower activities while showing limited improvement on the most dynamic ones.")
    add_figure(doc, "paper_threshold_validation/plots/adl_false_positive_by_task.png",
               "Figure 2. False-positive rate by activity-of-daily-living task, Rule A versus Rule B.")

    add_heading(doc, "5.3 Fainting-While-Sitting Subset", level=2)
    add_para(doc,
        "The source publication notes that acceleration-based detection is less reliable for falls "
        "caused by fainting while seated, recommending gyroscope-based detection instead. On this "
        "subset (461 trials), Rule A achieved 100.0% detection, Rule B achieved 87.4%, and Rule C "
        "(gyroscope-based) achieved 100.0%, confirming the gyroscope signal is a viable and reliable "
        "alternative for this fall subtype, consistent with the source publication's guidance. Rule C "
        "was not evaluated for specificity and its threshold was not independently tuned; this result "
        "should be read as confirming sensitivity only.")
    add_figure(doc, "paper_threshold_validation/plots/heatmap_ruleB.png",
               "Figure 3. Rule B detection rate by subject and fall task.")
    add_figure(doc, "paper_threshold_validation/plots/fainting_subset_axis_comparison.png",
               "Figure 4. Detection rate on the fainting-while-sitting subset, by rule.")

    add_heading(doc, "5.4 Phase-Aligned Signal Shape", level=2)
    add_para(doc,
        "The event-locked grand average across all 2,319 fall trials shows acceleration magnitude and "
        "vertical velocity each following a sharply-timed, low-variance trajectory: acceleration "
        "magnitude remains near a steady 1 g baseline, dips to approximately 0.6 g, and spikes to "
        "approximately 3.0 g precisely at the labeled impact frame; vertical velocity rises smoothly "
        "from near zero, crossing 0.3 m/s early in the fall and peaking near 2.2 m/s immediately before "
        "impact. Gyroscope magnitude, added in the four-signal analysis, shows an equally sharp and "
        "well-timed profile, rising smoothly from a low baseline to a pronounced peak at impact. "
        "Orientation tilt, by contrast, shows substantial variance even prior to fall onset and no "
        "sharply localized feature, indicating it is a comparatively unreliable standalone marker.")
    add_figure(doc, "analysis_4panel/plots/pattern_grand_average.png",
               "Figure 5. Event-locked grand average across all fall trials "
               "(acceleration magnitude, vertical velocity, tilt, gyroscope magnitude).")

    add_heading(doc, "5.5 Inter-Subject and Intra-Subject Consistency", level=2)
    add_para(doc,
        "Inter-subject analysis (holding fall type constant, comparing across the 32 subjects) shows "
        "individual subject mean traces closely bundled around the global mean for acceleration "
        "magnitude, vertical velocity, and gyroscope magnitude, with visibly wider dispersion for tilt. "
        "Intra-subject analysis (holding subject constant, comparing across each subject's own trials) "
        "shows a consistent, sharply-timed acceleration peak for every one of the 32 subjects "
        "individually. A trial-level raster of all 2,319 individual trials confirms the pattern is "
        "present at the single-trial level and is not an artifact of averaging.")
    add_figure(doc, "analysis_4panel/plots/pattern_intersubject_overlay.png",
               "Figure 6. Inter-subject overlay: each line represents one subject's mean trajectory.")
    add_figure(doc, "analysis_3panel/plots/pattern_intrasubject_small_multiples.png",
               "Figure 7. Intra-subject consistency: individual trials and subject mean, "
               "acceleration magnitude, per subject.")
    add_figure(doc, "analysis_3panel/plots/pattern_trial_heatmap.png",
               "Figure 8. Single-trial raster of acceleration magnitude across all 2,319 fall trials.")

    # 6. Discussion
    add_heading(doc, "6. Discussion", level=1)
    add_para(doc,
        "The evidence gathered across sensitivity/specificity testing and phase-aligned signal "
        "visualization is consistent and complementary. Three of the four raw signal channels "
        "available from a single wearable inertial sensor -- acceleration magnitude, vertical "
        "velocity, and gyroscope magnitude -- exhibit a fall signature that is sharply timed, low in "
        "cross-subject variance, and reproducible at the level of individual trials across the entire "
        "dataset. This satisfies the original question of whether the pattern is global: it is not "
        "specific to the individual example trial that motivated the investigation, nor to any single "
        "subject or fall type.")
    add_para(doc,
        "A distinct finding is that orientation tilt, while used as a confirmation signal in the "
        "source publication's benchmark algorithm, is markedly less consistent than the other three "
        "channels. This is consistent with, and explains, the comparatively lower detection rates "
        "observed for slow-onset fall types (forward-fall-while-sitting-down and forward-fall-while-"
        "sitting-caused-by-fainting) under Rule B.")
    add_para(doc,
        "Generalization of the underlying signal pattern should not be conflated with adequacy of a "
        "simple threshold rule as a deployable detector. Even the most complete threshold rule "
        "evaluated (Rule B) produces a substantial false-positive rate on ordinary dynamic activities. "
        "This is expected: the source publication's own comparative results show that trained "
        "classifiers (support vector machine, convolutional-recurrent deep learning) achieve "
        "substantially higher specificity (94.87% and 99.01% respectively) than its threshold-based "
        "algorithm (83.43%) on the same underlying signals.")

    # 7. Conclusion
    add_heading(doc, "7. Conclusion", level=1)
    add_para(doc,
        "The KFall dataset exhibits a genuine, generalizable pre-impact fall signature across "
        "subjects and fall types, carried primarily by acceleration magnitude, vertical velocity, and "
        "gyroscope magnitude. This supports the dataset's suitability for pre-impact fall detection "
        "research. Orientation tilt is a comparatively weak standalone signal. Fixed-threshold rules, "
        "including the source publication's own, are sufficient to demonstrate the pattern exists but "
        "are not sufficient as standalone detectors; a trained classifier operating on these same "
        "validated signals is the appropriate next step for a deployable system.")

    # 8. Recommendations
    add_heading(doc, "8. Recommendations", level=1)
    add_bullets(doc, [
        "Proceed to a trained classifier (support vector machine or deep learning) using the four "
        "validated signal channels as features, consistent with the source publication's own "
        "benchmark results.",
        "Re-evaluate specificity using a streaming, debounced detection protocol rather than a "
        "whole-recording match, for a more direct comparison against the source publication's reported "
        "figures.",
        "Independently calibrate the gyroscope-based rule (Rule C) and evaluate its specificity "
        "against activities of daily living before relying on it beyond the fainting-while-sitting "
        "subset.",
        "Treat orientation tilt as a secondary, not primary, feature in any subsequent detection "
        "model.",
    ])

    # 9. References
    add_heading(doc, "9. References", level=1)
    add_para(doc,
        "Yu, X., Jang, J., & Xiong, S. (2021). A Large-Scale Open Motion Dataset (KFall) and Benchmark "
        "Algorithms for Detecting Pre-impact Fall of the Elderly Using Wearable Inertial Sensors. "
        "Frontiers in Aging Neuroscience, 13:692865. https://doi.org/10.3389/fnagi.2021.692865")

    # 10. Appendix
    add_heading(doc, "10. Appendix: Implementation File Structure", level=1)
    add_table(doc,
        ["Path", "Contents"],
        [
            ["paper_threshold_validation/analyze_pattern.py", "Rule A/B/C signal computation and sensitivity/specificity testing (also the shared pipeline other folders import)"],
            ["paper_threshold_validation/pattern_results.csv", "Per-trial results underlying Section 5.1-5.3"],
            ["analysis_3panel/", "Phase-aligned visualization: acceleration, vertical velocity, tilt"],
            ["analysis_4panel/", "Phase-aligned visualization: adds gyroscope magnitude"],
            ["signal_quality/", "Interactive Plotly dashboards (subject/task explorer)"],
            ["rolling_regression/", "Causal Savitzky-Golay slope/curvature (beta1/beta2) signals"],
            ["ensemble_trigger/", "Voting-ensemble detector (methodology/docs stage)"],
            ["docs/pattern_analysis.md", "Technical implementation walkthrough (full project history)"],
            ["docs/paper.pdf", "Source publication"],
        ])

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    build()
